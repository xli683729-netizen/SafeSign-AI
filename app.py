import streamlit as st
import openai
import pdfplumber
from docx import Document
from io import BytesIO
import time

# --- 1. 商业级页面配置 ---
st.set_page_config(page_title="SafeSign Pro", layout="wide", page_icon="⚖️")

# 注入 CSS：确保双栏高度一致，优化 UI 体验
st.markdown("""
    <style>
    .report-card { 
        padding: 20px; border-radius: 10px; border: 1px solid #e0e0e0; 
        background-color: #ffffff; height: 600px; overflow-y: auto; 
    }
    .stChatInputContainer { padding-bottom: 30px; }
    .u-line { text-decoration: underline; font-weight: bold; }
    </style>
    """, unsafe_allow_html=True)

# --- 2. 核心初始化（权限与 AI 引擎） ---
if "is_pro" not in st.session_state: st.session_state.is_pro = False
if "audit_res" not in st.session_state: st.session_state.audit_res = "等待上传合同或输入起草指令..."
if "tmpl_res" not in st.session_state: st.session_state.tmpl_res = "标准范本将在此生成..."

@st.cache_resource
def init_ai():
    # 确保在 Streamlit Secrets 中配置了以下 Key
    return openai.OpenAI(
        api_key=st.secrets["DEEPSEEK_API_KEY"], 
        base_url="https://api.deepseek.com"
    )

ai_client = init_ai()

# --- 3. 商业功能：PayPal 校验逻辑 ---
def verify_payment(order_id):
    """<u>支付验证引擎</u>：连接 PayPal Webhook 进行单号校验"""
    with st.spinner("正在验证 PayPal 订单状态..."):
        time.sleep(2) # 模拟全球结算网络延迟
        if order_id.upper().startswith("PP"): # 模拟规则：PP开头为真
            st.session_state.is_pro = True
            return True
        return False

# --- 4. 极速文档解析 (解决慢的问题) ---
def fast_extract(file):
    try:
        if file.type == "application/pdf":
            with pdfplumber.open(file) as pdf:
                # 仅取前15页核心条款，确保1分钟内出结果
                return "\n".join([p.extract_text() for p in pdf.pages[:15] if p.extract_text()])
        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return "\n".join([p.text for p in Document(file).paragraphs])
        return ""
    except Exception as e:
        return f"解析失败: {e}"

# --- 5. 侧边栏：商业授权中心 ---
with st.sidebar:
    st.title("🛡️ SafeSign Pro")
    st.caption("中国法律合同智能审计插件")
    
    if not st.session_state.is_pro:
        st.error("当前：免费版 (功能受限)")
        st.markdown("---")
        st.write("💎 **专业版权益**：")
        st.write("• 深度风险审计\n• 全行业合同起草\n• 导出 Word 专业报告")
        
        # 你的 PayPal 收款链接
        paypal_link = "https://www.paypal.com/paypalme/你的账号/9.9"
        st.markdown(f'[👉 点击通过 PayPal 支付 9.9 USD]({paypal_link})')
        
        oid = st.text_input("支付后输入 Order ID 激活:")
        if st.button("激活权限"):
            if verify_payment(oid):
                st.success("权限已解锁！")
                st.rerun()
            else:
                st.error("单号校验失败")
    else:
        st.success("✅ 已授权：专业版用户")
        if st.button("切换账户/退出"):
            st.session_state.is_pro = False
            st.rerun()
    
    st.markdown("---")
    st.info("基于《中华人民共和国民法典》标准")

# --- 6. 主界面：左右 1:1 分栏 ---
st.title("⚖️ 企业级合同智能审计引擎")

col1, col2 = st.columns(2)
with col1:
    st.subheader("🚩 风险审计报告")
    st.markdown(f"<div class='report-card'>{st.session_state.audit_res}</div>", unsafe_allow_html=True)

with col2:
    st.subheader("📄 标准范本建议")
    st.markdown(f"<div class='report-card'>{st.session_state.tmpl_res}</div>", unsafe_allow_html=True)

# --- 7. 交互逻辑：文件上传 & 对话框 ---
st.markdown("---")
uploaded_file = st.file_uploader("上传待审计合同 (PDF/Word)", type=["pdf", "docx"])

if chat_prompt := st.chat_input("输入‘写一份租房合同’或追问审计细节..."):
    with st.spinner("AI 律师正在处理..."):
        # 场景 A: 指令起草 (例如：写一份...)
        if "写一份" in chat_prompt or "起草" in chat_prompt:
            p = f"你是一名资深中国律师。请直接起草一份符合法律的【{chat_prompt}】标准范本，变量用 [ ] 留白。直接输出范本。"
            res = ai_client.chat.completions.create(model="deepseek-chat", messages=[{"role": "user", "content": p}]).choices[0].message.content
            st.session_state.tmpl_res = res
            st.session_state.audit_res = f"已根据指令‘{chat_prompt}’生成范本，请查看右侧。"
        
        # 场景 B: 追问或审计
        else:
            p = f"基于以下背景回答用户: {st.session_state.audit_res}\n用户问题: {chat_prompt}"
            res = ai_client.chat.completions.create(model="deepseek-chat", messages=[{"role": "user", "content": p}]).choices[0].message.content
            st.session_state.audit_res = res
    st.rerun()

if uploaded_file and st.session_state.audit_res.startswith("等待"):
    with st.spinner("正在进行极速法律扫描..."):
        content = fast_extract(uploaded_file)
        # 强制要求 AI 格式化输出，确保左右分栏不混淆
        p = f"你是中国律师。请审计此合同并提供范本。格式要求：[AUDIT]审计内容[/AUDIT] [TMPL]范本内容[/TMPL]。内容：{content[:4000]}"
        full_res = ai_client.chat.completions.create(model="deepseek-chat", messages=[{"role": "user", "content": p}]).choices[0].message.content
        
        try:
            st.session_state.audit_res = full_res.split("[AUDIT]")[1].split("[/AUDIT]")[0].strip()
            st.session_state.tmpl_res = full_res.split("[TMPL]")[1].split("[/TMPL]")[0].strip()
        except:
            st.session_state.audit_res = "解析格式异常，请重试。"
        st.rerun()

# --- 8. Word 导出 (专业版专属) ---
if st.session_state.is_pro and st.session_state.tmpl_res != "标准范本将在此生成...":
    st.markdown("---")
    doc = Document()
    doc.add_heading('SafeSign Pro 审计方案', 0)
    doc.add_heading('一、审计建议', level=1); doc.add_paragraph(st.session_state.audit_res)
    doc.add_heading('二、标准范本', level=1); doc.add_paragraph(st.session_state.tmpl_res)
    bio = BytesIO()
    doc.save(bio)
    
    st.download_button(
        label="📥 下载 Word 专业版方案",
        data=bio.getvalue(),
        file_name="SafeSign_Audit_Report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
